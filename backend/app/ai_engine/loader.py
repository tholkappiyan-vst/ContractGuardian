"""Document loading: accepts raw bytes or file paths, returns LangChain Documents."""
import io
from pathlib import Path
from langchain_core.documents import Document
from langchain_community.document_loaders import PyMuPDFLoader
from docx import Document as DocxFile
from PIL import Image
import pytesseract
import fitz  # PyMuPDF


class ContractLoader:
    """Loads contracts from bytes or file paths into LangChain Documents."""

    SUPPORTED = {"pdf", "docx", "png", "jpg", "jpeg", "tiff"}

    def load_from_bytes(self, content: bytes, file_type: str, filename: str = "") -> list[Document]:
        file_type = file_type.lower().strip(".")
        if file_type not in self.SUPPORTED:
            raise ValueError(f"Unsupported: {file_type}. Supported: {self.SUPPORTED}")

        if file_type == "pdf":
            return self._load_pdf_bytes(content, filename)
        elif file_type == "docx":
            return self._load_docx_bytes(content, filename)
        else:
            return self._load_image_bytes(content, filename)

    def load_from_path(self, path: str | Path) -> list[Document]:
        path = Path(path)
        ext = path.suffix.lower().strip(".")
        content = path.read_bytes()
        return self.load_from_bytes(content, ext, path.name)

    def _load_pdf_bytes(self, content: bytes, filename: str) -> list[Document]:
        doc = fitz.open(stream=content, filetype="pdf")
        documents = []

        for i, page in enumerate(doc):
            text = page.get_text()

            # If no text extracted, OCR the page
            if not text.strip():
                pix = page.get_pixmap(dpi=300)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                text = pytesseract.image_to_string(img)

            if text.strip():
                documents.append(Document(
                    page_content=text,
                    metadata={
                        "source": filename,
                        "page": i + 1,
                        "total_pages": len(doc),
                        "file_type": "pdf",
                    },
                ))

        doc.close()
        return documents

    def _load_docx_bytes(self, content: bytes, filename: str) -> list[Document]:
        doc = DocxFile(io.BytesIO(content))
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)

        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    full_text.append(row_text)

        text = "\n\n".join(full_text)
        return [Document(
            page_content=text,
            metadata={"source": filename, "file_type": "docx"},
        )]

    def _load_image_bytes(self, content: bytes, filename: str) -> list[Document]:
        img = Image.open(io.BytesIO(content))
        text = pytesseract.image_to_string(img)

        if not text.strip():
            raise ValueError(f"OCR extracted no text from {filename}")

        return [Document(
            page_content=text,
            metadata={"source": filename, "file_type": "image", "ocr": True},
        )]
